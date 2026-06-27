#!/usr/bin/env python3
"""Basic QA for humanities-paper DOCX files."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from zipfile import ZipFile


def load_docx():
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        print(json.dumps({"error": f"python-docx is unavailable: {exc}"}, ensure_ascii=False))
        sys.exit(2)
    return Document


def iter_header_texts(document):
    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            text = "\n".join(p.text for p in header.paragraphs).strip()
            if text:
                yield text


def iter_footer_texts(document):
    for section in document.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            text = "\n".join(p.text for p in footer.paragraphs).strip()
            if text:
                yield text


def inspect_docx_package(path: Path):
    with ZipFile(path) as archive:
        names = archive.namelist()
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8") if "word/styles.xml" in names else ""
        rels_xml = (
            archive.read("word/_rels/document.xml.rels").decode("utf-8")
            if "word/_rels/document.xml.rels" in names
            else ""
        )
    blue_markers = {'w:val="0000FF"', 'w:val="0563C1"', 'w:val="1F4E79"', 'w:val="2F5496"'}
    return {
        "table_xml": "<w:tbl" in document_xml,
        "outline_level_count": document_xml.count("<w:outlineLvl"),
        "header_parts": [name for name in names if name.startswith("word/header")],
        "footer_parts": [name for name in names if name.startswith("word/footer")],
        "rels_has_header_footer": ("header" in rels_xml.lower()) or ("footer" in rels_xml.lower()),
        "blue_style_or_direct_xml": any(marker in document_xml or marker in styles_xml for marker in blue_markers),
    }


def count_blue_runs(document):
    count = 0
    samples = []
    for para in document.paragraphs:
        for run in para.runs:
            color = run.font.color.rgb
            if color is None:
                continue
            rgb = str(color).upper()
            if rgb.endswith("FF") or rgb in {"0000FF", "0563C1", "1F4E79", "2F5496"}:
                count += 1
                if len(samples) < 5 and run.text.strip():
                    samples.append(run.text.strip()[:40])
    return count, samples


def reference_section_stats(document):
    paragraphs = [p.text.strip() for p in document.paragraphs]
    start = None
    for index, text in enumerate(paragraphs):
        if text in {"参考文献", "References"}:
            start = index + 1
            break
    if start is None:
        return {
            "reference_heading_found": False,
            "reference_paragraphs": 0,
            "reference_single_block_risk": False,
            "reference_samples": [],
        }
    refs = [text for text in paragraphs[start:] if text]
    date_patterns = [len(re.findall(r"\(\d{4}[a-z]?\)", ref)) for ref in refs]
    return {
        "reference_heading_found": True,
        "reference_paragraphs": len(refs),
        "reference_single_block_risk": bool(len(refs) == 1 and date_patterns and date_patterns[0] >= 3),
        "reference_samples": refs[:3],
    }


def style_stats(document):
    names = [p.style.name for p in document.paragraphs]
    heading_style_paragraphs = [name for name in names if name.lower().startswith("heading")]
    return {
        "paragraph_styles": sorted(set(names)),
        "heading_style_paragraphs": len(heading_style_paragraphs),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Run basic structural QA on a DOCX file.")
    parser.add_argument("docx", type=Path)
    args = parser.parse_args()

    if not args.docx.exists():
        print(json.dumps({"error": f"file not found: {args.docx}"}, ensure_ascii=False))
        return 1

    Document = load_docx()
    document = Document(str(args.docx))
    text = "\n".join(p.text for p in document.paragraphs)
    headers = list(iter_header_texts(document))
    footers = list(iter_footer_texts(document))
    blue_runs, blue_samples = count_blue_runs(document)
    package = inspect_docx_package(args.docx)
    references = reference_section_stats(document)
    styles = style_stats(document)
    result = {
        "file": str(args.docx),
        "paragraphs": len(document.paragraphs),
        "paragraph_styles": styles["paragraph_styles"],
        "heading_style_paragraphs": styles["heading_style_paragraphs"],
        "outline_level_count": package["outline_level_count"],
        "tables": len(document.tables),
        "table_xml": package["table_xml"],
        "headers_nonempty": len(headers),
        "header_samples": headers[:5],
        "footers_nonempty": len(footers),
        "footer_samples": footers[:5],
        "header_parts": package["header_parts"],
        "footer_parts": package["footer_parts"],
        "rels_has_header_footer": package["rels_has_header_footer"],
        "chinese_chars": len(re.findall(r"[\u4e00-\u9fff]", text)),
        "nonspace_chars": len(re.sub(r"\s+", "", text)),
        "benwen_count": text.count("本文"),
        "ascii_double_quotes": text.count('"'),
        "chinese_double_quotes": text.count("“") + text.count("”"),
        "fullwidth_colons": text.count("："),
        "not_but_patterns": len(re.findall(r"不是[^。；]*而是", text)),
        "not_but_patterns_bingfei": len(re.findall(r"并非[^。；]*而是", text)),
        "reference_heading_found": references["reference_heading_found"],
        "reference_paragraphs": references["reference_paragraphs"],
        "reference_single_block_risk": references["reference_single_block_risk"],
        "reference_samples": references["reference_samples"],
        "blue_runs": blue_runs,
        "blue_samples": blue_samples,
        "blue_style_or_direct_xml": package["blue_style_or_direct_xml"],
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
